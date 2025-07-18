from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import user_passes_test
from django.http import JsonResponse
from .models import FileUpload
import os
from django.conf import settings
from django.contrib.auth.decorators import login_required, user_passes_test
from .models import Source_file, Sources,Notes
from django.contrib.auth import authenticate, login, logout
from django.contrib import messages
from collections import defaultdict
from django.shortcuts import render, get_object_or_404

def index(request):
    return render (request, 'index.html')

# Admin Login View (without using Django forms)
def admin_login(request):
    if request.user.is_authenticated and request.user.is_superuser:
        return redirect('upload_pdf')  # Redirect if already logged in
    
    if request.method == "POST":
        username = request.POST.get('username')
        password = request.POST.get('password')

        user = authenticate(request, username=username, password=password)

        if user is not None and user.is_superuser:
            login(request, user)
            return redirect('index')  # Redirect to upload page
        else:
            messages.error(request, "Invalid Credentials or Not a Superuser")
    
    return render(request, 'admin/admin_login.html')

# Admin Logout View
@login_required
def admin_logout(request):
    logout(request)
    return redirect('admin_login')

# Check if user is a superuser
def superuser_required(user):
    return user.is_superuser

# View to upload PDF (only for superusers)
@login_required
@user_passes_test(superuser_required)
def upload_pdf(request):
    if request.method == 'POST':
        course_id = request.POST.get('course')
        pdf_file = request.FILES.get('file')
        file_title = request.POST.get('title')  # Get title from form input

        if course_id and pdf_file:
            course = Sources.objects.get(id=course_id)
            Source_file.objects.create(course=course, file=pdf_file, title=file_title)
            return redirect('view_pdfs')

    courses = Sources.objects.all()
    return render(request, 'admin/upload_pdf.html', {'courses': courses})

from django.shortcuts import render, get_object_or_404
from .models import Sources, Source_file

#import os
import mimetypes
from django.shortcuts import render, get_object_or_404
from django.http import FileResponse, Http404
from docx import Document as DocxDocument
from .models import Sources, Source_file

# View to list all courses
def course_list(request):
    courses = Sources.objects.all()
    return render(request, "viewer/viewer_base.html", {"courses": courses})

# View to list PDFs/DOCs for a specific course
def pdf_list(request, course_id):
    course = get_object_or_404(Sources, id=course_id)
    
    # Fetch Source Files
    pdf_files = Source_file.objects.filter(course=course)
    
    # Fetch Notes categorized by type
    class_notes = Notes.objects.filter(course=course, note_type='class')
    test_notes = Notes.objects.filter(course=course, note_type='test')
    infographic_notes = Notes.objects.filter(course=course, note_type='infographic')

    # Group by title
    grouped_files = defaultdict(list)
    grouped_class_notes = defaultdict(list)
    grouped_test_notes = defaultdict(list)
    grouped_infographic_notes = defaultdict(list)
    print(infographic_notes)

    for file in pdf_files:
        grouped_files[file.title].append(file)
    for note in class_notes:
        grouped_class_notes[note.title].append(note)
    for note in test_notes:
        grouped_test_notes[note.title].append(note)
    for note in infographic_notes:
        grouped_infographic_notes[note.title].append(note)

    return render(request, "admin/pdf_view.html", {
        "course": course,
        "grouped_files": dict(grouped_files),
        "grouped_class_notes": dict(grouped_class_notes),
        "grouped_test_notes": dict(grouped_test_notes),
        "grouped_infographic_notes": dict(grouped_infographic_notes),
    })



def notes_list(request, course_id, note_type):
    course = get_object_or_404(Sources, id=course_id)
    notes = Notes.objects.filter(course=course, note_type=note_type)
    return render(request, "students/notes_list.html", {"course": course, "notes": notes, "note_type": note_type})

# Function to read DOCX files
def read_docx(file_path):
    try:
        doc = DocxDocument(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        return text
    except Exception as e:
        return "Error reading file."

# View to display a specific PDF/DOCX file
from django.shortcuts import render, get_object_or_404
from django.http import HttpResponse
from docx import Document as DocxDocument


import os
import mimetypes
from django.shortcuts import render, get_object_or_404
from django.http import HttpResponse
from docx import Document as DocxDocument
from .models import Source_file, Notes  # Import Notes model

def pdf_detail(request, pdf_id):
    # Try to get from Source_file first
    document = Source_file.objects.filter(id=pdf_id).first()

    # If not found in Source_file, check in Notes
    if not document:
        document = get_object_or_404(Notes, id=pdf_id)  

    file_path = document.file.path
    file_type, _ = mimetypes.guess_type(file_path)

    if not os.path.exists(file_path):
        return HttpResponse("File not found.", status=404)

    # Convert DOCX to paginated text
    content = []
    if file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = DocxDocument(file_path)
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        content = [paragraphs[i:i+10] for i in range(0, len(paragraphs), 10)]  # 10 lines per page

    return render(request, "admin/pdf_detail.html", {
        "document": document,
        "file_type": file_type,
        "file_url": document.file.url,
        "content": content,  # Send paginated DOCX text
        "total_pages": len(content),
    })
import docx
from django.shortcuts import render, redirect
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from exam.models import Question, Course

@login_required
def upload_questions_docx(request):
    if not request.user.is_superuser:
        messages.error(request, "Access denied! Only superusers can upload files.")
        return redirect("home")

    if request.method == "POST" and request.FILES.get("file"):
        file = request.FILES["file"]

        try:
            doc = docx.Document(file)
            table = doc.tables[0]  # Assumes first table contains data

            headers = [cell.text.strip() for cell in table.rows[0].cells]
            for row in table.rows[1:]:
                data = {headers[i]: cell.text.strip() for i, cell in enumerate(row.cells)}

                try:
                    course = Course.objects.get(course_name=data["Course"])
                    Question.objects.update_or_create(
                        course=course,
                        question=data["Question"],
                        defaults={
                            "marks": int(data["Marks"]),
                            "priority": data["Priority"],
                            "option1": data["Option1"],
                            "option2": data["Option2"],
                            "option3": data["Option3"],
                            "option4": data["Option4"],
                            "answer": data["Answer"],
                            "reason": data.get("Reason", "No explanation provided."),
                        }
                    )
                except Course.DoesNotExist:
                    messages.error(request, f"Course not found: {data['Course']}")
                    continue

            messages.success(request, "Questions uploaded successfully!")
            return redirect("upload_questions_docx")

        except Exception as e:
            messages.error(request, f"Error processing file: {str(e)}")

    return render(request, "exam/upload_questions_docx.html")
